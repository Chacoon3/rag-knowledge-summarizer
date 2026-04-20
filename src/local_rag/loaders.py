from __future__ import annotations

import asyncio
from pathlib import Path
from typing import Iterable

from docx import Document as DocxDocument
from pypdf import PdfReader

from local_rag.models import LoadedDocument


SUPPORTED_SUFFIXES = {".txt", ".md", ".pdf", ".docx"}
TEXT_ENCODINGS = ("utf-8", "utf-8-sig", "gb18030")


def discover_files(source_dir: Path) -> list[Path]:
    return sorted(
        path
        for path in source_dir.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_SUFFIXES
    )


def load_documents(source_dir: Path) -> list[LoadedDocument]:
    source_dir = Path(source_dir)
    if not source_dir.exists():
        raise FileNotFoundError(f"文档目录不存在: {source_dir}")

    return load_documents_from_paths(discover_files(source_dir), source_root=source_dir)


async def load_documents_async(source_dir: Path) -> list[LoadedDocument]:
    source_dir = Path(source_dir)
    if not source_dir.exists():
        raise FileNotFoundError(f"文档目录不存在: {source_dir}")

    paths = await discover_files_async(source_dir)
    return await load_documents_from_paths_async(paths, source_root=source_dir)


def load_documents_from_paths(
    paths: Iterable[Path],
    source_root: Path | None = None,
) -> list[LoadedDocument]:
    if source_root is not None:
        source_root = Path(source_root)

    documents: list[LoadedDocument] = []
    for raw_path in paths:
        path = Path(raw_path)
        if not path.exists():
            raise FileNotFoundError(f"文档文件不存在: {path}")
        if not path.is_file():
            continue
        if path.suffix.lower() not in SUPPORTED_SUFFIXES:
            raise ValueError(f"不支持的文件类型: {path.suffix or path.name}")

        content = extract_text(path).strip()
        if not content:
            continue

        source_path = path.name
        if source_root is not None:
            try:
                source_path = path.relative_to(source_root).as_posix()
            except ValueError:
                source_path = path.name

        documents.append(
            LoadedDocument(
                source_path=source_path,
                content=content,
            )
        )
    return documents


async def discover_files_async(source_dir: Path) -> list[Path]:
    return await asyncio.to_thread(discover_files, source_dir)


async def load_documents_from_paths_async(
    paths: Iterable[Path],
    source_root: Path | None = None,
) -> list[LoadedDocument]:
    if source_root is not None:
        source_root = Path(source_root)

    normalized_paths: list[Path] = []
    for raw_path in paths:
        path = Path(raw_path)
        if not path.exists():
            raise FileNotFoundError(f"文档文件不存在: {path}")
        if not path.is_file():
            continue
        if path.suffix.lower() not in SUPPORTED_SUFFIXES:
            raise ValueError(f"不支持的文件类型: {path.suffix or path.name}")
        normalized_paths.append(path)

    tasks = [
        _load_document_from_path_async(path, source_root=source_root)
        for path in normalized_paths
    ]
    results = await asyncio.gather(*tasks)
    return [document for document in results if document is not None]


async def _load_document_from_path_async(
    path: Path,
    source_root: Path | None = None,
) -> LoadedDocument | None:
    content = (await extract_text_async(path)).strip()
    if not content:
        return None

    source_path = path.name
    if source_root is not None:
        try:
            source_path = path.relative_to(source_root).as_posix()
        except ValueError:
            source_path = path.name

    return LoadedDocument(source_path=source_path, content=content)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return read_text_file(path)
    if suffix == ".pdf":
        return read_pdf_file(path)
    if suffix == ".docx":
        return read_docx_file(path)
    raise ValueError(f"不支持的文件类型: {path.suffix}")


async def extract_text_async(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return await read_text_file_async(path)
    if suffix == ".pdf":
        return await read_pdf_file_async(path)
    if suffix == ".docx":
        return await read_docx_file_async(path)
    raise ValueError(f"不支持的文件类型: {path.suffix}")


def read_text_file(path: Path) -> str:
    for encoding in TEXT_ENCODINGS:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


async def read_text_file_async(path: Path) -> str:
    for encoding in TEXT_ENCODINGS:
        try:
            return await asyncio.to_thread(path.read_text, encoding=encoding)
        except UnicodeDecodeError:
            continue
    return await asyncio.to_thread(path.read_text, encoding="utf-8", errors="ignore")


def read_pdf_file(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


async def read_pdf_file_async(path: Path) -> str:
    return await asyncio.to_thread(read_pdf_file, path)


def read_docx_file(path: Path) -> str:
    document = DocxDocument(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


async def read_docx_file_async(path: Path) -> str:
    return await asyncio.to_thread(read_docx_file, path)
